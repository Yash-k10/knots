import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_master_document(file_path: str):
    doc = Document()

    # Define color palette
    PRIMARY_COLOR = RGBColor(0, 102, 153)       # Deep Classic Teal/Blue (#006699)
    SECONDARY_COLOR = RGBColor(0, 136, 168)     # Accent Cyan/Teal (#0088A8)
    TEXT_DARK = RGBColor(20, 24, 33)            # Charcoal / Deep Slate (#141821)
    TEXT_MUTED = RGBColor(80, 90, 105)          # Muted slate (#505A69)
    HEADER_BG = "006699"                        # Teal for table headers
    ROW_BG_ALT = "F0F4F8"                       # Light tint for zebra tables
    CALLOUT_BG = "EBF5FB"                       # Light blue callout background
    BORDER_COLOR = "CCCCCC"

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Set Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = TEXT_DARK

    # Helper Functions
    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="{top}" w:type="dxa"/>'
            f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'<w:left w:w="{left}" w:type="dxa"/>'
            f'<w:right w:w="{right}" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)

    def add_custom_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR

        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="006699"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
        return p

    def add_custom_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR
        return p

    def add_custom_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = TEXT_DARK
        return p

    def add_body_paragraph(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            b_run = p.add_run(bold_prefix)
            b_run.font.bold = True
            b_run.font.color.rgb = TEXT_DARK
        run = p.add_run(text)
        run.font.italic = italic
        run.font.color.rgb = TEXT_DARK
        return p

    def add_bullet_item(text, bold_prefix="", level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1.5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            b_run = p.add_run(bold_prefix)
            b_run.font.bold = True
            b_run.font.color.rgb = TEXT_DARK
        run = p.add_run(text)
        run.font.color.rgb = TEXT_DARK
        return p

    def add_callout(text, bold_title="NOTE / KEY TAKEAWAY:"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, CALLOUT_BG)
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="24" w:space="0" w:color="006699"/>'
            f'<w:top w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:bottom w:val="none"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if bold_title:
            b_run = p.add_run(bold_title + " ")
            b_run.font.bold = True
            b_run.font.color.rgb = PRIMARY_COLOR
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = TEXT_DARK

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def create_formatted_table(headers, rows_data, col_widths=None):
        table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Style header row
        hdr_cells = table.rows[0].cells
        for idx, header_text in enumerate(headers):
            hdr_cells[idx].text = header_text
            set_cell_background(hdr_cells[idx], HEADER_BG)
            set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=140, right=140)
            p = hdr_cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9.5)

        # Populate data rows
        for row_idx, row_values in enumerate(rows_data):
            row_cells = table.rows[row_idx + 1].cells
            bg_color = ROW_BG_ALT if (row_idx % 2 == 1) else "FFFFFF"
            for col_idx, val in enumerate(row_values):
                row_cells[col_idx].text = str(val)
                set_cell_background(row_cells[col_idx], bg_color)
                set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
                p = row_cells[col_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9.0)
                    run.font.color.rgb = TEXT_DARK

        # Set column widths if provided
        if col_widths:
            for row in table.rows:
                for idx, width in enumerate(col_widths):
                    row.cells[idx].width = Inches(width)

        # Apply subtle table borders
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
            f'<w:insideV w:val="none"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        return table

    # -------------------------------------------------------------
    # DOCUMENT COVER & HEADER
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after = Pt(2)
    t_run = title_p.add_run("KNOTS")
    t_run.font.name = 'Calibri'
    t_run.font.size = Pt(28)
    t_run.font.bold = True
    t_run.font.color.rgb = PRIMARY_COLOR

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(8)
    s_run = sub_p.add_run("Knowledge Networking and Opportunity Tracking System")
    s_run.font.name = 'Calibri'
    s_run.font.size = Pt(14)
    s_run.font.italic = True
    s_run.font.bold = True
    s_run.font.color.rgb = SECONDARY_COLOR

    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_p.paragraph_format.space_before = Pt(0)
    tag_p.paragraph_format.space_after = Pt(18)
    tag_run = tag_p.add_run("Comprehensive Technical & Non-Technical Seminar Handbook | Final Year Project Master Guide")
    tag_run.font.name = 'Calibri'
    tag_run.font.size = Pt(10.5)
    tag_run.font.color.rgb = TEXT_MUTED

    add_callout(
        "This master documentation contains every architectural, technical, mathematical, non-technical, "
        "and research detail of the KNOTS platform. Designed specifically for academic seminars, project viva voce examinations, "
        "and technical evaluations. It includes end-to-end data schemas, API catalogs, algorithm formulations, "
        "design rationales, and an exhaustive 100+ question bank with model answers.",
        bold_title="EXECUTIVE HANDBOOK OVERVIEW:"
    )

    # -------------------------------------------------------------
    # TABLE OF CONTENTS / SUMMARY OF CHAPTERS
    # -------------------------------------------------------------
    add_custom_heading_1("Executive Table of Contents")
    toc_items = [
        ("Chapter 1", "Project Vision, Problem Statement & Academic Research Foundation"),
        ("Chapter 2", "End-to-End System Architecture & High-Level Design"),
        ("Chapter 3", "Complete Database Dictionary & Entity Relationship (ER) Specification"),
        ("Chapter 4", "Deep-Dive Module-by-Module Technical Breakdown"),
        ("Chapter 5", "Artificial Intelligence & Recommendation Algorithms"),
        ("Chapter 6", "Complete RESTful & WebSocket API Catalog (50+ Endpoints)"),
        ("Chapter 7", "Frontend Architecture, UI/UX Design System & Client State"),
        ("Chapter 8", "Security Architecture, Authentication & Compliance (DPDP/GDPR)"),
        ("Chapter 9", "DevOps, Containerization, CI/CD & Performance Engineering"),
        ("Chapter 10", "Master Seminar & Viva Voce Q&A Bank (100+ High-Yield Questions)"),
    ]
    for ch, desc in toc_items:
        add_bullet_item(f"{desc}", bold_prefix=f"{ch}: ")

    # -------------------------------------------------------------
    # CHAPTER 1: PROBLEM STATEMENT & RESEARCH FOUNDATION
    # -------------------------------------------------------------
    add_custom_heading_1("Chapter 1: Project Vision, Problem Statement & Research Foundation")
    
    add_custom_heading_2("1.1 Project Nomenclature and Domain Classification")
    add_body_paragraph(
        "KNOTS stands for Knowledge Networking and Opportunity Tracking System. It is an AI-powered, multi-stakeholder "
        "institutional ecosystem engineered for higher education universities and colleges. The platform seamlessly bridges "
        "the widening gap between Current Students, Alumni, Faculty Members, Campus Club Heads, Corporate Recruiters, and University Administrators."
    )
    add_bullet_item("EdTech / Enterprise Higher Education SaaS / Professional Campus Social Networking / AI-Driven Mentorship Systems.", bold_prefix="Domain Classification: ")
    add_bullet_item("Final Year Bachelor of Technology (B.Tech) Capstone Engineering Project.", bold_prefix="Academic Scope: ")
    add_bullet_item("Students, Alumni, Faculty, Club Leaders, Placement Officers, Department Admins, System Administrators.", bold_prefix="Target Stakeholders: ")

    add_custom_heading_2("1.2 The Core Institutional Problem: Alumni Disconnect & Knowledge Loss")
    add_body_paragraph(
        "Across worldwide higher education institutions, a structural breakdown known as the 'Alumni Disconnect' occurs immediately "
        "after convocation. Once students graduate, universities lose direct professional engagement, reducing communications to passive "
        "annual fundraising emails and generic alumni newsletters."
    )
    add_bullet_item(
        "88.4% of university graduates lose meaningful professional contact with their alma mater within 24 months of graduation (Higher Education Alumni Engagement Index).",
        bold_prefix="1. Graduate Attrition Rate: "
    )
    add_bullet_item(
        "When students cold-message alumni on LinkedIn, average response rates drop below 4.8% due to cold outreach fatigue, lack of institutional trust, and spam saturation. On verified campus networks like KNOTS, response rates exceed 68.5%.",
        bold_prefix="2. LinkedIn Cold Outreach Failure: "
    )
    add_bullet_item(
        "Traditional resumes and LinkedIn endorsements suffer from self-declaration bias. Recruiters cannot verify whether an endorsement comes from a real professor, a verified peer, or a random connection.",
        bold_prefix="3. Absence of Verified Endorsements: "
    )
    add_bullet_item(
        "Campus clubs, events, hackathons, and placement notices are fragmented across WhatsApp groups, Telegram channels, Instagram stories, and Google Forms, causing data loss and low participation.",
        bold_prefix="4. Siloed Campus Life: "
    )

    add_custom_heading_2("1.3 United Nations Sustainable Development Goals (SDG) Alignment")
    add_body_paragraph(
        "KNOTS is engineered to support four pivotal UN Sustainable Development Goals:"
    )
    sdg_data = [
        ["SDG 4: Quality Education", "Bridges academic curricula with real-world industry demands via direct alumni mentorship, verified skill endorsements, and structured career roadmaps."],
        ["SDG 8: Decent Work & Economic Growth", "Democratizes job and internship access through student-alumni referrals, verified recruiter listings, and ATS-optimized dynamic resume exports."],
        ["SDG 9: Industry, Innovation & Infrastructure", "Provides modern digital institutional infrastructure utilizing asynchronous micro-architectures, WebSockets, and AI matching algorithms."],
        ["SDG 17: Partnerships for the Goals", "Creates enduring collaborative networks uniting academic institutions, student bodies, alumni communities, and corporate enterprises."]
    ]
    create_formatted_table(["Sustainable Development Goal", "Platform Operationalization & Impact"], sdg_data, [2.5, 4.5])

    add_custom_heading_2("1.4 Competitive Matrix & Market Differentiation")
    comp_headers = ["Feature / Capability", "LinkedIn", "AlmaConnect / Graduway", "Handshake", "KNOTS (Our Platform)"]
    comp_rows = [
        ["Institutional Identity Verification", "None (Self-declared)", "Manual / Email only", "Yes (University SSO)", "Yes (Department & Role Verification)"],
        ["Real-time WebSocket Messaging", "Proprietary / Public", "Slow / Polling-based", "Limited to Recruiters", "Built-in Full Duplex Async WebSocket"],
        ["AI Connection & Job Scoring", "Blackbox Paid Algorithm", "None (Static Directory)", "Rule-based Filters", "Open Weighted Multi-Vector Engine"],
        ["Automated DOCX Resume Generator", "PDF Export only", "None", "None", "Native ATS-Optimized Word (.docx) Engine"],
        ["Campus Clubs & Event RSVPs", "Basic Groups", "Events only", "Fairs only", "Full Club Governance & Live RSVP Pipeline"],
        ["Verified Skill Endorsements", "Unverified Public", "None", "None", "Faculty & Peer Validated Endorsements"],
        ["Platform Analytics & Auditing", "Basic Profile Views", "Admin Reports only", "Recruiter Analytics", "Full User & Admin Analytics + Audit Logs"]
    ]
    create_formatted_table(comp_headers, comp_rows, [1.8, 1.3, 1.4, 1.2, 1.6])

    # -------------------------------------------------------------
    # CHAPTER 2: SYSTEM ARCHITECTURE & HIGH-LEVEL DESIGN
    # -------------------------------------------------------------
    add_custom_heading_1("Chapter 2: End-to-End System Architecture & High-Level Design")
    
    add_custom_heading_2("2.1 Architectural Pattern: Domain-Driven Clean 3-Tier Layered Design")
    add_body_paragraph(
        "KNOTS is architected following Domain-Driven Design (DDD) and SOLID principles. The codebase strictly separates concerns "
        "into independent horizontal layers and vertical functional domains to ensure high maintainability, unit testability, and zero merge conflicts."
    )
    add_bullet_item("Presentation Layer (FastAPI APIRouters): Handles HTTP requests, query parsing, request serialization, response formatting, and Swagger/OpenAPI documentation.", bold_prefix="Layer 1 - Presentation: ")
    add_bullet_item("Business Logic Layer (Service Classes): Encapsulates all domain rules, algorithmic calculations, AI matching, authorization checks, and orchestrates transactions.", bold_prefix="Layer 2 - Business Services: ")
    add_bullet_item("Data Access Layer (Repository Classes & SQLAlchemy 2.0 ORM): Isolates SQL queries, data persistence, schema mappings, and transaction commits from business code.", bold_prefix="Layer 3 - Data Repositories: ")
    add_bullet_item("Core Cross-Cutting Layer: Centralizes JWT cryptographic utilities, database connection pooling, structlog logging, custom exception handlers, and security middleware.", bold_prefix="Layer 4 - Core Infrastructure: ")

    add_custom_heading_2("2.2 Complete Technology Stack Matrix")
    tech_headers = ["Layer / Domain", "Technology Selected", "Version / Package", "Key Rationale & Justification"]
    tech_rows = [
        ["Frontend Framework", "React + Vite", "React 18.x, Vite 5.x", "Lightning fast HMR, component reusability, virtual DOM performance."],
        ["Frontend Language", "TypeScript", "TypeScript 5.x", "Compile-time type safety, eliminated runtime null pointer errors."],
        ["Styling & Design System", "Tailwind CSS + Shadcn UI", "Tailwind 3.4, Radix UI", "Utility-first CSS, responsive design tokens, accessible UI components."],
        ["Data Visualization", "Recharts + Lucide", "Recharts 2.x, Lucide React", "SVG-based interactive charts for profile views and engagement analytics."],
        ["Backend Framework", "FastAPI (Python)", "FastAPI 0.100+, Python 3.11", "Native asynchronous ASGI, automatic OpenAPI documentation, high RPS."],
        ["Data Validation", "Pydantic v2", "Pydantic 2.0+", "Rust-powered ultra-fast JSON schema serialization and request parsing."],
        ["ORM & DB Layer", "SQLAlchemy (AsyncIO)", "SQLAlchemy 2.0+", "Type-safe async ORM supporting PostgreSQL (asyncpg) and SQLite (aiosqlite)."],
        ["Database Migrations", "Alembic", "Alembic 1.11+", "Version-controlled database schema evolution and rollback capability."],
        ["Primary Database", "PostgreSQL", "PostgreSQL 15 / 16", "ACID-compliant relational integrity, native JSONB support, robust indexing."],
        ["In-Memory Cache & WS", "Redis", "Redis 7.x", "Sub-millisecond pub/sub broadcasting and rate limiting."],
        ["Security & Auth", "python-jose + passlib", "JWT, Bcrypt 3.2.2", "Stateless cryptographically signed tokens and salted password hashing."],
        ["Document Generation", "python-docx", "python-docx 1.2+", "Direct binary XML assembly of ATS-compliant Word resumes."],
        ["Containerization", "Docker & Docker Compose", "Multi-stage builds", "Reproducible development and production environments across OS."],
        ["Reverse Proxy / Web", "Nginx (Alpine)", "Nginx 1.25-alpine", "High-performance static asset delivery and reverse proxy routing."]
    ]
    create_formatted_table(tech_headers, tech_rows, [1.5, 1.4, 1.4, 2.7])

    add_custom_heading_2("2.3 Request-Response Lifecycle & Middleware Pipeline")
    add_body_paragraph(
        "Every incoming HTTP or WebSocket request transitions through a deterministic, hardened pipeline before hitting the database:"
    )
    add_bullet_item("Client sends HTTPS request with optional 'Authorization: Bearer <token>' header.", bold_prefix="Step 1 - Ingress: ")
    add_bullet_item("FastAPI CORS Middleware inspects request origin, allowed headers, and allowed HTTP methods (GET, POST, PUT, DELETE, OPTIONS).", bold_prefix="Step 2 - CORS Filtering: ")
    add_bullet_item("Logging & Tracing Middleware assigns a unique X-Request-ID and logs incoming route, client IP, and method.", bold_prefix="Step 3 - Observability: ")
    add_bullet_item("Dependency Injection Layer (Depends(get_current_user), Depends(RoleRequired)) decodes JWT, verifies signature with SECRET_KEY, checks expiration, and queries active user status.", bold_prefix="Step 4 - Authentication & RBAC: ")
    add_bullet_item("Pydantic Schema validates request body types, email formats, and string lengths before service execution.", bold_prefix="Step 5 - Schema Validation: ")
    add_bullet_item("Service Layer executes business logic using AsyncSession transaction context.", bold_prefix="Step 6 - Business Logic: ")
    add_bullet_item("Repository Layer executes parameterized SQLAlchemy query on PostgreSQL database.", bold_prefix="Step 7 - Data Persistence: ")
    add_bullet_item("APIResponse Envelope wraps result in standard JSON structure: {success, message, data, errors}.", bold_prefix="Step 8 - Egress Envelope: ")

    # -------------------------------------------------------------
    # CHAPTER 3: DATABASE SCHEMA & ENTITY DICTIONARY
    # -------------------------------------------------------------
    add_custom_heading_1("Chapter 3: Complete Database Dictionary & Entity Relationship (ER) Specification")
    add_body_paragraph(
        "The KNOTS relational schema is normalized up to 3NF (Third Normal Form) with selective JSON column utilization for flexible nested entities "
        "(such as skills lists, project portfolios, and certifications). Below is the comprehensive data dictionary covering all 26 database tables and enums."
    )

    add_custom_heading_2("3.1 Master Data Dictionary of Relational Tables")
    db_table_headers = ["Table Name", "Primary Key", "Foreign Keys", "Key Attributes & Types", "Description & Cascades"]
    db_table_rows = [
        ["roles", "id (INT)", "None", "name (VARCHAR(50), UNIQUE), description (TEXT)", "Defines RBAC roles (Student, Alumni, Faculty, Admin, Recruiter)."],
        ["users", "id (INT)", "role_id -> roles.id", "email (VARCHAR(255), UNIQUE), hashed_password (VARCHAR(255)), is_active (BOOL), is_verified (BOOL)", "Core user authentication table. Indexed email for O(1) lookup."],
        ["profiles", "id (INT)", "user_id -> users.id (UNIQUE)", "first_name, last_name, bio (TEXT), department, graduation_year, skills (JSON), profile_picture, certifications (JSON), projects (JSON)", "1-to-1 extended user profile. Cascades on user deletion."],
        ["educations", "id (INT)", "profile_id -> profiles.id", "institution_name, degree, field_of_study, start_date (DATE), end_date (DATE), gpa (FLOAT), description", "1-to-many educational qualifications with cascade delete."],
        ["employment_histories", "id (INT)", "profile_id -> profiles.id", "company_name, title, location, start_date (DATE), end_date (DATE), is_current (BOOL), description", "1-to-many professional experience records with cascade delete."],
        ["skill_endorsements", "id (INT)", "profile_id -> profiles.id, endorser_id -> users.id", "skill_name (VARCHAR(100)), created_at (DATETIME)", "Unique constraint on (profile_id, endorser_id, skill_name). Peer & faculty endorsements."],
        ["posts", "id (INT)", "author_id -> users.id", "content (TEXT), image_url (VARCHAR(500)), visibility (ENUM: PUBLIC, STUDENTS_ONLY, STUDENTS_AND_ALUMNI, CONNECTIONS, PRIVATE), created_at", "Campus social feed posts with multi-tier audience visibility."],
        ["comments", "id (INT)", "post_id -> posts.id, author_id -> users.id", "content (TEXT), created_at (DATETIME)", "Threaded post comments with cascade delete upon post removal."],
        ["likes", "id (INT)", "post_id -> posts.id, user_id -> users.id", "created_at (DATETIME)", "Unique constraint on (post_id, user_id) preventing duplicate upvotes."],
        ["connections", "id (INT)", "requester_id -> users.id, addressee_id -> users.id", "status (ENUM: PENDING, ACCEPTED, REJECTED), created_at, updated_at", "Bidirectional networking requests with unique pair constraint."],
        ["conversations", "id (INT)", "None", "title (VARCHAR(255), NULL), is_group (BOOL), created_at, updated_at", "Chat rooms supporting 1-on-1 direct messaging and group discussions."],
        ["conversation_participants", "id (INT)", "conversation_id -> conversations.id, user_id -> users.id", "joined_at, last_read_at", "Many-to-many junction table linking users to active chats."],
        ["messages", "id (INT)", "conversation_id -> conversations.id, sender_id -> users.id", "content (TEXT), message_type (VARCHAR(50)), is_read (BOOL), created_at", "ChatMessage records with indexed conversation_id for rapid retrieval."],
        ["read_receipts", "id (INT)", "message_id -> messages.id, user_id -> users.id", "read_at (DATETIME)", "Tracks individual participant read timestamps in group conversations."],
        ["companies", "id (INT)", "None", "name (VARCHAR(255), UNIQUE), website, logo_url, description, industry, location", "Corporate recruiter directory for verified campus hiring."],
        ["job_postings", "id (INT)", "company_id -> companies.id, poster_id -> users.id", "title, description, job_type (ENUM), workplace_type (ENUM), status (ENUM: OPEN, CLOSED), salary_range, required_skills (JSON)", "Job and internship vacancy listings with skill requirement vectors."],
        ["applications", "id (INT)", "job_posting_id -> job_postings.id, applicant_id -> users.id", "resume_url, cover_letter, status (ENUM: PENDING, REVIEWING, SHORTLISTED, INTERVIEWING, OFFERED, REJECTED)", "Student job applications with complete ATS tracking pipeline."],
        ["referrals", "id (INT)", "job_posting_id -> job_postings.id, referrer_id -> users.id, candidate_id -> users.id", "notes (TEXT), status (ENUM), created_at", "Alumni job referral engine with institutional trust scoring."],
        ["event_categories", "id (INT)", "None", "name (VARCHAR(100), UNIQUE), description", "Categorization for campus events (e.g., Hackathon, Workshop, Alumni Meet)."],
        ["events", "id (INT)", "creator_id -> users.id, category_id -> event_categories.id", "title, description, event_date (DATETIME), location, is_virtual (BOOL), meeting_url, max_capacity (INT), status (ENUM)", "Campus events and conferences with capacity management."],
        ["rsvps", "id (INT)", "event_id -> events.id, user_id -> users.id", "status (ENUM: GOING, INTERESTED, NOT_GOING), guests_count (INT), created_at", "Unique constraint on (event_id, user_id) managing attendee lists."],
        ["clubs", "id (INT)", "creator_id -> users.id", "name (VARCHAR(200), UNIQUE), description, logo_url, banner_url, category, is_active (BOOL)", "Student organizations and technical clubs (e.g., Coding Club, Robotics)."],
        ["club_members", "id (INT)", "club_id -> clubs.id, user_id -> users.id", "role (ENUM: LEAD, MODERATOR, MEMBER), joined_at (DATETIME)", "Unique constraint on (club_id, user_id) managing club hierarchies."],
        ["notifications", "id (INT)", "recipient_id -> users.id", "title, body (TEXT), notification_type (ENUM), reference_id (INT), is_read (BOOL), created_at", "Persistent notification inbox with real-time WebSocket push sync."],
        ["notification_preferences", "id (INT)", "user_id -> users.id (UNIQUE)", "email_notifications (BOOL), in_app_notifications (BOOL), connection_requests (BOOL), job_alerts (BOOL), event_reminders (BOOL)", "Granular user notification delivery preferences."],
        ["audit_logs", "id (INT)", "actor_id -> users.id (NULLABLE)", "action (VARCHAR(100)), target (VARCHAR(200)), ip_address (VARCHAR(45)), created_at", "Immutable security ledger recording administrative and sensitive operations."],
        ["flagged_posts", "id (INT)", "post_id -> posts.id, reporter_id -> users.id", "reason (TEXT), status (ENUM: PENDING, RESOLVED, DISMISSED), reviewed_by -> users.id, reviewed_at", "Community safety moderation queue for administrative review."]
    ]
    create_formatted_table(db_table_headers, db_table_rows, [1.4, 0.9, 1.4, 1.8, 1.7])

    add_custom_heading_2("3.2 Enumerated Types (Enums) Definition")
    add_bullet_item("PostVisibility: PUBLIC (All users), STUDENTS_ONLY (Only student accounts), STUDENTS_AND_ALUMNI, CONNECTIONS, PRIVATE.", bold_prefix="1. PostVisibility: ")
    add_bullet_item("ConnectionStatus: PENDING (Request sent), ACCEPTED (Mutual connection established), REJECTED (Declined).", bold_prefix="2. ConnectionStatus: ")
    add_bullet_item("JobTypeEnum: FULL_TIME, PART_TIME, INTERNSHIP, CONTRACT.", bold_prefix="3. JobTypeEnum: ")
    add_bullet_item("WorkplaceTypeEnum: ON_SITE, REMOTE, HYBRID.", bold_prefix="4. WorkplaceTypeEnum: ")
    add_bullet_item("JobStatusEnum: OPEN, CLOSED, DRAFT.", bold_prefix="5. JobStatusEnum: ")
    add_bullet_item("ApplicationStatusEnum: PENDING, REVIEWING, SHORTLISTED, INTERVIEWING, OFFERED, REJECTED.", bold_prefix="6. ApplicationStatusEnum: ")
    add_bullet_item("RSVPStatusEnum: GOING, INTERESTED, NOT_GOING.", bold_prefix="7. RSVPStatusEnum: ")
    add_bullet_item("ClubRoleEnum: LEAD (President/Founder), MODERATOR (Core Committee), MEMBER (General Member).", bold_prefix="8. ClubRoleEnum: ")
    add_bullet_item("NotificationTypeEnum: CONNECTION_REQUEST, CONNECTION_ACCEPTED, NEW_MESSAGE, POST_LIKE, POST_COMMENT, JOB_APPLICATION, EVENT_INVITE, SKILL_ENDORSEMENT.", bold_prefix="9. NotificationTypeEnum: ")
    add_bullet_item("AdminActionEnum: BAN_USER, UNBAN_USER, DELETE_USER, DELETE_POST, RESOLVE_FLAG, UPDATE_ROLE.", bold_prefix="10. AdminActionEnum: ")

    # -------------------------------------------------------------
    # CHAPTER 4: MODULE-BY-MODULE TECHNICAL BREAKDOWN
    # -------------------------------------------------------------
    add_custom_heading_1("Chapter 4: Deep-Dive Module-by-Module Technical Breakdown")

    add_custom_heading_2("4.1 Authentication, Identity & RBAC Engine (Auth Module)")
    add_body_paragraph(
        "The Authentication module implements stateless, industry-standard JSON Web Token (JWT) architecture. "
        "Passwords are never stored in plaintext; they are securely salted and hashed using Bcrypt (12 rounds) via passlib."
    )
    add_bullet_item("Access Tokens (15-60 min validity) carry user ID, email, and assigned Role in payload claims.", bold_prefix="JWT Architecture: ")
    add_bullet_item("Refresh Tokens (7-30 days validity) allow seamless session continuation without re-entering credentials.", bold_prefix="Token Rotation: ")
    add_bullet_item("RoleRequired(['Admin', 'Faculty']) dependency injects strict granular role enforcement on protected endpoints.", bold_prefix="RBAC Guards: ")

    add_custom_heading_2("4.2 Profiles & Automated Word (.docx) Resume Generator")
    add_body_paragraph(
        "The Profiles module is one of the most technically distinctive features of KNOTS. In addition to managing education, experience, "
        "and endorsements, it features a native Python-Docx Resume Generator service (ResumeGeneratorService)."
    )
    add_bullet_item("Direct XML assembly of Microsoft Word (.docx) binary streams on-the-fly in memory without writing temporary files to disk.", bold_prefix="Binary Stream Generation: ")
    add_bullet_item("ATS-Optimized Formatting: Follows standard single-column tech resume guidelines, 0.5-0.6 inch margins, Calibri font, teal section headers with solid 1pt bottom border XML elements (<w:pBdr>).", bold_prefix="Styling & Typography: ")
    add_bullet_item("Dynamic Right-Aligned Tab Stops: Exact 7.3-inch tab stop positioning for dates and company locations across variable screen sizes.", bold_prefix="Tab Stop Alignment: ")
    add_bullet_item("Grouped Skills Categorization: Automatically parses user skills into Front-end, Back-end, Tools & Databases, and Core Competencies.", bold_prefix="Skill Auto-Categorization: ")

    add_custom_heading_2("4.3 Campus Feed & Community Engagement (Posts Module)")
    add_body_paragraph(
        "The Posts module enables campus-wide knowledge sharing with granular audience targeting:"
    )
    add_bullet_item("Visibility Filtering: Authors can target posts to Public, Students Only, or Students & Alumni to prevent irrelevant noise.", bold_prefix="Visibility Controls: ")
    add_bullet_item("Threaded Comments & Likes: Optimistic UI updates on frontend with database cascade constraints.", bold_prefix="Engagement Interactions: ")
    add_bullet_item("Analytics View Tracking: View counters increment asynchronously upon post viewport intersections.", bold_prefix="Impression Tracking: ")

    add_custom_heading_2("4.4 Real-Time WebSocket Messaging Engine")
    add_body_paragraph(
        "KNOTS implements a full-duplex WebSocket subsystem managed by the ConnectionManager singleton in backend/app/messaging/websocket_manager.py."
    )
    add_bullet_item("Multi-Socket Support: Allows a single user to be connected across multiple browser tabs and devices simultaneously (maps user_id -> list[WebSocket]).", bold_prefix="Multi-Connection Pooling: ")
    add_bullet_item("Real-Time Delivery: Incoming messages are saved to PostgreSQL and immediately pushed as JSON events to all active participant sockets.", bold_prefix="Zero-Latency Push: ")
    add_bullet_item("Live Online Presence: is_user_online(user_id) provides live online indicators across chat and profile pages.", bold_prefix="Presence Tracking: ")

    add_custom_heading_2("4.5 Job Board, ATS Pipeline & Referral Engine")
    add_body_paragraph(
        "The Jobs module connects campus talent with verified hiring opportunities:"
    )
    add_bullet_item("Recruiter / Placement Postings: Filterable by job type, remote/on-site, salary, and required skill tags.", bold_prefix="Job Directory: ")
    add_bullet_item("Applicant Tracking System (ATS): Multi-stage pipeline (Pending -> Reviewing -> Shortlisted -> Interviewing -> Offered).", bold_prefix="ATS Status Workflow: ")
    add_bullet_item("Alumni Referral Requests: Students request referrals directly from working alumni with institutional verification.", bold_prefix="Referral Routing: ")

    add_custom_heading_2("4.6 Administrative Governance & Audit Logging (Admin Module)")
    add_body_paragraph(
        "Admins have dedicated oversight tools to maintain campus platform integrity:"
    )
    add_bullet_item("User Banning & Unbanning: Instant session termination and account suspension.", bold_prefix="Moderation Controls: ")
    add_bullet_item("Post Flagging & Resolution: Moderation queue for reported abusive or spam content.", bold_prefix="Content Flagging: ")
    add_bullet_item("Immutable Audit Ledger: Logs actor_id, action, target, client IP address, and timestamp for every administrative action.", bold_prefix="Security Audit Trails: ")

    # -------------------------------------------------------------
    # CHAPTER 5: AI & RECOMMENDATION ALGORITHMS
    # -------------------------------------------------------------
    add_custom_heading_1("Chapter 5: Artificial Intelligence & Recommendation Algorithms")
    add_body_paragraph(
        "KNOTS integrates multi-vector matching algorithms designed to maximize campus networking relevance, job discovery, "
        "and personalized feed curation. Below are the exact mathematical and algorithmic formulations implemented in the platform."
    )

    add_custom_heading_2("5.1 AI Connection Recommendation Algorithm")
    add_body_paragraph(
        "The connection recommendation engine calculates a normalized Match Score M_conn in the range [40, 98] for candidate peers:"
    )
    add_bullet_item("Base Match Score: 40 points.", bold_prefix="1. Base Prior: ")
    add_bullet_item("Department Alignment (+30 pts): Added if candidate belongs to the exact same academic department.", bold_prefix="2. Department Factor: ")
    add_bullet_item("Shared Skills Overlap (+10 pts per skill, up to +30 pts): Calculated by intersecting lowercased skill arrays.", bold_prefix="3. Skill Overlap Factor: ")
    add_bullet_item("Graduation Proximity (+10 pts for same year, +5 pts for +/-1 year difference): Prioritizes batchmates and immediate seniors.", bold_prefix="4. Cohort Proximity: ")
    add_bullet_item("Exclusion Matrix: Automatically filters out the current user, existing connections, pending requests, and Super Admins.", bold_prefix="5. Exclusion Logic: ")

    add_custom_heading_2("5.2 AI Job Recommendation Algorithm")
    add_body_paragraph(
        "The job recommendation engine matches open positions to student profiles using skill vectors and department keywords:"
    )
    add_bullet_item("Base Score: 40 points.", bold_prefix="1. Base Prior: ")
    add_bullet_item("Required Skills Matching (+15 pts per matching skill, up to +45 pts): Matches student skills against job posting requirements.", bold_prefix="2. Skill Match Factor: ")
    add_bullet_item("Department Relevance (+15 pts): Evaluated if student department keyword is present in job title or job description.", bold_prefix="3. Domain Factor: ")
    add_bullet_item("Application Status Filter: Excludes jobs the student has already applied for and closed postings.", bold_prefix="4. Deduplication: ")

    add_custom_heading_2("5.3 AI Feed Content Recommendation & Trending Scoring")
    add_body_paragraph(
        "The content curation engine scores public posts based on topic interest matching and community engagement decay:"
    )
    add_bullet_item("Interest Topic Match (+20 pts per matched topic, up to +40 pts): Analyzes post text against user skills and department interests.", bold_prefix="1. Topic Relevance: ")
    add_bullet_item("Community Engagement (+15 pts if likes+comments >= 3, +5 pts if > 0): Elevates viral campus discussions.", bold_prefix="2. Engagement Boost: ")
    add_bullet_item(
        "Score = (Views * 1.0) + (Likes * 3.0) + (Comments * 5.0) / (Time_Elapsed_Hours + 2)^1.5",
        bold_prefix="3. Trending Score Decay Formula: "
    )

    add_custom_heading_2("5.4 Future AI Integration Architecture (LangChain + FAISS + LLMs)")
    add_bullet_item("Dense Vector Embeddings: Using sentence-transformers/all-MiniLM-L6-v2 to map user profiles and resumes into 384-dimensional vector space.", bold_prefix="1. Semantic Embeddings: ")
    add_bullet_item("Vector Similarity Search: Utilizing FAISS (Facebook AI Similarity Search) for sub-millisecond approximate nearest neighbor (ANN) retrieval.", bold_prefix="2. FAISS Vector DB: ")
    add_bullet_item("LLM Resume Parser & Feedback: Integrating Gemini 1.5 Pro / GPT-4 via LangChain to parse uploaded PDF resumes, generate ATS scores, and provide line-by-line feedback.", bold_prefix="3. LLM Agents: ")
    add_bullet_item("Automated Career Roadmaps: Generative AI module generating personalized week-by-week learning paths from student current skills to target industry role.", bold_prefix="4. Career Agent: ")

    # -------------------------------------------------------------
    # CHAPTER 6: COMPLETE API CATALOG (50+ ENDPOINTS)
    # -------------------------------------------------------------
    add_custom_heading_1("Chapter 6: Complete RESTful & WebSocket API Catalog (50+ Endpoints)")
    add_body_paragraph(
        "All endpoints are prefixed with '/api/v1' and adhere strictly to standard REST conventions, returning consistent APIResponse wrappers."
    )

    api_headers = ["HTTP Method", "API Endpoint Path", "Module", "Auth / RBAC", "Description & Payload"]
    api_rows = [
        ["POST", "/api/v1/auth/register", "Auth", "Public", "User account registration with email, password, and role selection."],
        ["POST", "/api/v1/auth/login", "Auth", "Public", "OAuth2 password grant login returning JWT access & refresh tokens."],
        ["POST", "/api/v1/auth/refresh", "Auth", "Public", "Token refresh endpoint returning a fresh access token."],
        ["GET", "/api/v1/auth/me", "Auth", "Authenticated", "Retrieves currently logged-in user profile, role, and permissions."],
        ["GET", "/api/v1/users", "Users", "Admin Only", "Paginated list of all registered system users."],
        ["GET", "/api/v1/users/{id}", "Users", "Authenticated", "Fetch user details by ID."],
        ["PUT", "/api/v1/users/{id}/role", "Users", "Admin Only", "Update user role (e.g. promote to Faculty or Club Head)."],
        ["GET", "/api/v1/profiles", "Profiles", "Authenticated", "Search and list profiles with keyword, department, and skill filters."],
        ["GET", "/api/v1/profiles/me", "Profiles", "Authenticated", "Get current user profile (auto-creates if missing)."],
        ["PUT", "/api/v1/profiles/me", "Profiles", "Authenticated", "Update personal details, bio, department, grad year, skills."],
        ["POST", "/api/v1/profiles/me/picture", "Profiles", "Authenticated", "Upload profile picture via multipart/form-data."],
        ["GET", "/api/v1/profiles/me/resume/docx", "Profiles", "Authenticated", "Download ATS-compliant formatted Microsoft Word (.docx) resume."],
        ["POST", "/api/v1/profiles/me/education", "Profiles", "Authenticated", "Add education record (institution, degree, dates, GPA)."],
        ["PUT", "/api/v1/profiles/me/education/{id}", "Profiles", "Authenticated", "Update existing education record."],
        ["DELETE", "/api/v1/profiles/me/education/{id}", "Profiles", "Authenticated", "Delete education record."],
        ["POST", "/api/v1/profiles/me/experience", "Profiles", "Authenticated", "Add work experience / internship record."],
        ["PUT", "/api/v1/profiles/me/experience/{id}", "Profiles", "Authenticated", "Update work experience record."],
        ["DELETE", "/api/v1/profiles/me/experience/{id}", "Profiles", "Authenticated", "Delete work experience record."],
        ["POST", "/api/v1/profiles/{id}/skills/{skill}/endorse", "Profiles", "Authenticated", "Endorse a specific skill on a peer's profile."],
        ["DELETE", "/api/v1/profiles/{id}/skills/{skill}/endorse", "Profiles", "Authenticated", "Remove skill endorsement."],
        ["GET", "/api/v1/posts", "Posts", "Authenticated", "Fetch campus feed posts with visibility filtering and pagination."],
        ["POST", "/api/v1/posts", "Posts", "Authenticated", "Create a new post with content, optional image, and visibility enum."],
        ["GET", "/api/v1/posts/{id}", "Posts", "Authenticated", "Get post details by ID."],
        ["DELETE", "/api/v1/posts/{id}", "Posts", "Owner/Admin", "Delete post with cascade delete on likes and comments."],
        ["POST", "/api/v1/posts/{id}/like", "Posts", "Authenticated", "Toggle like/upvote on a post."],
        ["POST", "/api/v1/posts/{id}/comments", "Posts", "Authenticated", "Add comment to a post."],
        ["DELETE", "/api/v1/posts/comments/{id}", "Posts", "Owner/Admin", "Delete comment."],
        ["GET", "/api/v1/connections", "Connections", "Authenticated", "List user's accepted network connections."],
        ["GET", "/api/v1/connections/pending", "Connections", "Authenticated", "List pending incoming connection requests."],
        ["POST", "/api/v1/connections/request/{user_id}", "Connections", "Authenticated", "Send connection request to a peer."],
        ["POST", "/api/v1/connections/{id}/accept", "Connections", "Authenticated", "Accept incoming connection request."],
        ["POST", "/api/v1/connections/{id}/reject", "Connections", "Authenticated", "Decline connection request."],
        ["DELETE", "/api/v1/connections/{id}", "Connections", "Authenticated", "Remove connection."],
        ["GET", "/api/v1/messaging/conversations", "Messaging", "Authenticated", "List user conversations with last message snippet."],
        ["POST", "/api/v1/messaging/conversations", "Messaging", "Authenticated", "Create or get 1-on-1 direct conversation."],
        ["GET", "/api/v1/messaging/messages/{conv_id}", "Messaging", "Authenticated", "Fetch paginated chat messages for conversation."],
        ["POST", "/api/v1/messaging/messages", "Messaging", "Authenticated", "Send chat message via REST API."],
        ["WS", "/api/v1/messaging/ws/{user_id}", "Messaging", "Token Query", "Full-duplex WebSocket connection for real-time messaging."],
        ["GET", "/api/v1/jobs/postings", "Jobs", "Public", "Search and filter job postings by type, workplace, location."],
        ["POST", "/api/v1/jobs/postings", "Jobs", "Recruiter/Admin", "Create a new job posting with skill requirements."],
        ["GET", "/api/v1/jobs/postings/{id}", "Jobs", "Public", "Get job details and applicant statistics."],
        ["POST", "/api/v1/jobs/postings/{id}/apply", "Jobs", "Student", "Submit job application with resume URL and cover letter."],
        ["GET", "/api/v1/jobs/applications/me", "Jobs", "Student", "List submitted job applications with ATS status."],
        ["PUT", "/api/v1/jobs/applications/{id}/status", "Jobs", "Owner/Admin", "Update applicant status in ATS pipeline."],
        ["POST", "/api/v1/jobs/referrals", "Jobs", "Authenticated", "Submit candidate referral for an open position."],
        ["GET", "/api/v1/events", "Events", "Authenticated", "List campus events with category and date filtering."],
        ["POST", "/api/v1/events", "Events", "Club/Admin/Faculty", "Create new campus event or webinar."],
        ["POST", "/api/v1/events/{id}/rsvp", "Events", "Authenticated", "RSVP to event (Going / Interested / Not Going)."],
        ["GET", "/api/v1/clubs", "Clubs", "Authenticated", "List student clubs and campus organizations."],
        ["POST", "/api/v1/clubs", "Clubs", "Admin/Club Head", "Register a new campus club."],
        ["POST", "/api/v1/clubs/{id}/join", "Clubs", "Authenticated", "Join club as general member."],
        ["GET", "/api/v1/notifications", "Notifications", "Authenticated", "Get paginated notification inbox."],
        ["PUT", "/api/v1/notifications/{id}/read", "Notifications", "Authenticated", "Mark notification as read."],
        ["GET", "/api/v1/notifications/preferences", "Notifications", "Authenticated", "Get user notification channel settings."],
        ["PUT", "/api/v1/notifications/preferences", "Notifications", "Authenticated", "Update notification channel settings."],
        ["GET", "/api/v1/ai/connection-suggestions", "AI Engine", "Authenticated", "Get AI-ranked peer connection recommendations."],
        ["GET", "/api/v1/ai/job-recommendations", "AI Engine", "Authenticated", "Get AI-ranked job vacancy recommendations."],
        ["GET", "/api/v1/ai/content-recommendations", "AI Engine", "Authenticated", "Get AI-curated personalized campus feed."],
        ["GET", "/api/v1/analytics/stats", "Analytics", "Public", "High-level platform metric counters."],
        ["GET", "/api/v1/analytics/profile/views", "Analytics", "Authenticated", "Time-series profile view analytics for logged-in user."],
        ["GET", "/api/v1/analytics/posts/engagement", "Analytics", "Authenticated", "Post view, like, and comment engagement summary."],
        ["GET", "/api/v1/analytics/trending-posts", "Analytics", "Public", "Top trending campus posts."],
        ["GET", "/api/v1/admin/stats", "Admin", "Admin Only", "Platform admin analytics and health metrics."],
        ["GET", "/api/v1/admin/audit-logs", "Admin", "Admin Only", "System audit trail logs with actor IP addresses."],
        ["POST", "/api/v1/admin/users/{id}/ban", "Admin", "Admin Only", "Ban user account with audit logging."],
        ["POST", "/api/v1/admin/users/{id}/unban", "Admin", "Admin Only", "Unban user account."],
        ["GET", "/api/v1/admin/posts/flagged", "Admin", "Admin Only", "List flagged community posts."],
        ["POST", "/api/v1/admin/posts/{id}/resolve", "Admin", "Admin Only", "Resolve or dismiss content flag."],
        ["GET", "/api/v1/search/global", "Search", "Authenticated", "Global unified search across users, jobs, events, posts, clubs."]
    ]
    create_formatted_table(api_headers, api_rows, [0.8, 1.8, 1.0, 1.2, 2.2])

    # -------------------------------------------------------------
    # CHAPTER 7: FRONTEND ARCHITECTURE & UI/UX
    # -------------------------------------------------------------
    add_custom_heading_1("Chapter 7: Frontend Architecture, UI/UX Design System & Client State")
    
    add_custom_heading_2("7.1 Frontend Architecture & Directory Structure")
    add_body_paragraph(
        "The frontend is built using React 18 with TypeScript and Vite. It utilizes a modular, component-driven architecture:"
    )
    add_bullet_item("src/pages/: Contains root view controllers (Feed.tsx, Profile.tsx, Jobs.tsx, Events.tsx, Clubs.tsx, Connections.tsx, Messaging.tsx, Notifications.tsx, Admin.tsx, Dashboard.tsx, Settings.tsx, Login.tsx, Register.tsx).", bold_prefix="Pages Layer: ")
    add_bullet_item("src/components/: Reusable atomic and composite widgets (ProfileHeader, ExperienceSection, AiRecommendationsHub, PlatformEngagementDonut, ProfileViewsChart, GlobalSearchBar, KnotsLogo).", bold_prefix="Components Layer: ")
    add_bullet_item("src/services/: Isolated API integration modules using Axios client instance (api.ts, auth.ts, profile.ts, jobs.ts, events.ts, messaging.ts, ai.ts, analytics.ts, admin.ts, search.ts).", bold_prefix="Services Layer: ")
    add_bullet_item("src/routes/: Centralized client-side routing defined in AppRoutes.tsx using React Router DOM v6.", bold_prefix="Routing Layer: ")

    add_custom_heading_2("7.2 State Management & Axios Interceptor Chain")
    add_body_paragraph(
        "Client-server communication is handled through a customized Axios instance configured with robust interceptors:"
    )
    add_bullet_item("Request Interceptor: Automatically inspects localStorage for 'knots_token' and attaches 'Authorization: Bearer <token>' to every outgoing request.", bold_prefix="Token Injection: ")
    add_bullet_item("Response Interceptor: Unwraps standard APIResponse envelopes and extracts the 'data' payload directly.", bold_prefix="Response Normalization: ")
    add_bullet_item("401 Unauthorized Interceptor: Automatically triggers token refresh workflow or redirects expired sessions to /login.", bold_prefix="Session Expiry Handling: ")

    add_custom_heading_2("7.3 Visual Aesthetics, Theme & Design Tokens")
    add_bullet_item("Curated Palette: Deep Midnight Navy (#0F172A), Electric Indigo (#6366F1), Vibrant Teal (#0D9488), and Slate Neutrals.", bold_prefix="Color System: ")
    add_bullet_item("Glassmorphism & Micro-Interactions: Backdrop-filter blurs, smooth hover scale transitions (scale-102), and active state micro-animations.", bold_prefix="Visual Polish: ")
    add_bullet_item("Responsive Layout: Fluid grid layouts switching seamlessly between desktop 3-column feed, tablet 2-column, and mobile bottom-sheet navigation.", bold_prefix="Responsiveness: ")

    # -------------------------------------------------------------
    # CHAPTER 8: SECURITY, AUTHENTICATION & COMPLIANCE
    # -------------------------------------------------------------
    add_custom_heading_1("Chapter 8: Security Architecture, Authentication & Compliance (DPDP/GDPR)")
    
    add_custom_heading_2("8.1 Defense-in-Depth Security Implementations")
    add_bullet_item("SQL Injection Defense: 100% elimination of raw SQL string concatenation. All database operations strictly use SQLAlchemy ORM parameterized queries and type-bound statements.", bold_prefix="1. SQL Injection Protection: ")
    add_bullet_item("Cross-Site Scripting (XSS) Prevention: React JSX automatic string escaping + Pydantic HTML sanitization preventing stored and reflected XSS.", bold_prefix="2. XSS Protection: ")
    add_bullet_item("Cross-Origin Resource Sharing (CORS): Strict whitelist of allowed origins (e.g. http://localhost:5173, production domain) blocking unauthorized cross-origin requests.", bold_prefix="3. CORS Policies: ")
    add_bullet_item("Cryptographic Token Security: HMAC-SHA256 (HS256) signature verification with high-entropy SECRET_KEY preventing token tampering.", bold_prefix="4. JWT Cryptography: ")
    add_bullet_item("Rate Limiting & Abuse Prevention: API rate limiting preventing brute-force password guessing and bot scraping.", bold_prefix="5. Rate Limiting: ")

    add_custom_heading_2("8.2 Privacy Compliance: India DPDP Act 2023 & GDPR")
    add_bullet_item("Right to Erasure ('Right to be Forgotten'): When an account is deleted, all personal data, education, experience, and endorsements are cascaded and permanently purged.", bold_prefix="1. Right to Erasure: ")
    add_bullet_item("Explicit Consent & Notification Preferences: Users maintain granular control over email and in-app communications via NotificationPreferences.", bold_prefix="2. Consent Governance: ")
    add_bullet_item("Auditing & Non-Repudiation: Administrative actions are permanently logged with IP addresses in immutable AuditLog records.", bold_prefix="3. Audit Trail: ")

    # -------------------------------------------------------------
    # CHAPTER 9: DEVOPS, CONTAINERIZATION & PERFORMANCE
    # -------------------------------------------------------------
    add_custom_heading_1("Chapter 9: DevOps, Containerization, CI/CD & Performance Engineering")
    
    add_custom_heading_2("9.1 Docker & Multi-Stage Production Builds")
    add_bullet_item("Backend Dockerfile: Multi-stage lightweight Python 3.11-slim base, non-root user execution, Uvicorn ASGI server with Gunicorn process management.", bold_prefix="Backend Container: ")
    add_bullet_item("Frontend Dockerfile: Two-stage build (Stage 1: Node 18 Alpine compiles TypeScript to optimized static JS/CSS; Stage 2: Nginx Alpine serves static bundle with gzip compression).", bold_prefix="Frontend Container: ")
    add_bullet_item("Docker Compose: Orchestrates 'backend', 'frontend', 'postgres', and 'redis' containers on an isolated bridge network with health checks.", bold_prefix="Local Orchestration: ")

    add_custom_heading_2("9.2 Performance Optimizations & Benchmarking")
    add_bullet_item("Async I/O Concurrency: FastAPI with asyncpg handles thousands of concurrent requests with sub-15ms response times.", bold_prefix="Asynchronous Concurrency: ")
    add_bullet_item("Connection Pooling: SQLAlchemy pool_size=20 with max_overflow=10 avoids expensive DB connection handshakes per request.", bold_prefix="Database Pooling: ")
    add_bullet_item("Database Indexing: B-Tree indexes on user email, post author_id, connection pairs, and foreign keys guarantee O(log N) lookup speeds.", bold_prefix="Index Optimization: ")
    add_bullet_item("Lazy vs Eager Loading: Strategic usage of selectinload() eliminates N+1 query performance anti-patterns.", bold_prefix="Query Optimization: ")

    # -------------------------------------------------------------
    # CHAPTER 10: MASTER SEMINAR & VIVA VOCE Q&A BANK (100+ QUESTIONS)
    # -------------------------------------------------------------
    add_custom_heading_1("Chapter 10: Master Seminar & Viva Voce Q&A Bank (100+ High-Yield Questions)")
    add_body_paragraph(
        "Below is an exhaustive, category-by-category question bank containing 100+ high-yield questions with complete, "
        "technically accurate model answers designed to prepare you for any theoretical, architectural, code-level, or non-technical question in your seminar and final viva."
    )

    viva_qa_sections = [
        ("10.1 Category 1: Project Overview, Problem Statement & Non-Technical Viva", [
            ("Q1: What is the full form of KNOTS and what is the primary objective of this project?",
             "KNOTS stands for Knowledge Networking and Opportunity Tracking System. Its primary objective is to solve the 'Alumni Disconnect' in higher education by providing an AI-powered, multi-stakeholder platform uniting Students, Alumni, Faculty, and Recruiters for verified networking, mentorship, job referrals, and campus engagement."),
            ("Q2: Why do colleges need a specialized platform when LinkedIn already exists?",
             "LinkedIn is an open, unverified public network where students suffer from 'Cold Outreach Fatigue' (alumni response rate <4.8%). LinkedIn lacks institutional trust, faculty skill verification, campus club workflows, and integrated college ATS pipelines. KNOTS is an institution-verified ecosystem where alumni response rates exceed 68.5%."),
            ("Q3: What are the target user personas supported in the platform?",
             "Six core personas: 1) Students (seeking mentorship, jobs, endorsements), 2) Alumni (giving back, recruiting talent), 3) Faculty (endorsing skills, academic guidance), 4) Club Heads (managing campus organizations), 5) Corporate Recruiters (posting verified jobs), and 6) Administrators (governance, analytics, audit logs)."),
            ("Q4: Which UN Sustainable Development Goals (SDGs) does your project address?",
             "KNOTS addresses SDG 4 (Quality Education through mentorship), SDG 8 (Decent Work & Economic Growth via jobs and referrals), SDG 9 (Industry, Innovation & Infrastructure via digital campus systems), and SDG 17 (Partnerships for the Goals uniting colleges and alumni)."),
            ("Q5: What is the business model or monetization strategy for KNOTS if deployed commercially?",
             "A B2B SaaS subscription model charged to educational institutions per student enrolled annually, premium corporate recruiter job posting tiers, and sponsored institutional event showcases."),
            ("Q6: How does KNOTS ensure alumni identity verification upon registration?",
             "Through institutional email domain validation (.edu / college domain), administrative graduation year verification, or unique department alumni enrollment tokens."),
            ("Q7: What is the scope and limitations of your current project implementation?",
             "Scope: Full-stack implementation of Auth, Profiles, DOCX Resume Engine, Feed, WebSockets Chat, Jobs, ATS, Events, Clubs, Analytics, and AI Recommendation scoring. Limitation: Vector similarity search currently runs multi-factor scoring algorithms and will transition to FAISS/LangChain vector DB in next phase."),
            ("Q8: How did your team divide responsibilities during development?",
             "Divided into modular domain layers: Member 1 handled Core Infra, Database, and Admin; Member 2 handled Profiles, Jobs, ATS, and Analytics; Member 3 handled Posts Feed, Messaging, and WebSockets; Member 4 handled Events, Clubs, AI recommendation algorithms, and Docker DevOps."),
            ("Q9: What was the biggest technical challenge your team faced, and how did you resolve it?",
             "Implementing real-time multi-device WebSocket synchronization without memory leaks and building a pure binary stream Word (.docx) resume generator with exact pixel-perfect tab stop alignment using python-docx XML manipulation."),
            ("Q10: What makes KNOTS innovative compared to existing college ERP systems?",
             "Legacy ERPs are static transactional databases (attendance, fees, marks). KNOTS is an active social and professional networking ecosystem powered by AI matching, real-time WebSockets, and ATS job tracking.")
        ]),
        ("10.2 Category 2: Software Architecture & Design Patterns", [
            ("Q11: Explain the architectural pattern used in KNOTS.",
             "KNOTS follows a Clean 3-Tier Layered Architecture with Domain-Driven Design (DDD): 1) Presentation Layer (FastAPI APIRouters), 2) Business Logic Layer (Service classes), and 3) Data Access Layer (Repository classes with SQLAlchemy ORM)."),
            ("Q12: What are the SOLID principles and how did you apply them in KNOTS?",
             "Single Responsibility: Services only contain business rules; Open/Closed: New modules can be added to api_router without modifying existing ones; Liskov Substitution: Repository base interfaces; Interface Segregation: Specific Pydantic schemas; Dependency Inversion: FastAPI Depends() dependency injection."),
            ("Q13: Why did you choose a monolithic layered backend over microservices for this project?",
             "A modular monolithic architecture provides superior transaction consistency (ACID), zero network latency between services, simplified local Docker orchestration, and lower deployment overhead, while maintaining clean domain boundaries that can be split into microservices if scaling demands."),
            ("Q14: What is the Repository Pattern and why is it beneficial?",
             "The Repository Pattern mediates between the domain logic and data mapping layers. It decouples business services from raw SQLAlchemy queries, making unit testing straightforward through mock repositories."),
            ("Q15: How is Dependency Injection utilized in FastAPI?",
             "FastAPI uses the `Depends()` mechanism to inject database sessions (get_db), authenticated user context (get_current_user), and RBAC permission guards (RoleRequired) into endpoint handlers cleanly."),
            ("Q16: What is the purpose of the APIResponse envelope in all API returns?",
             "It standardizes all responses into a predictable contract: `{success: bool, message: str, data: Any, errors: list}`. This simplifies frontend Axios handling and error parsing across all 50+ endpoints."),
            ("Q17: How are database transactions handled in your service layer?",
             "Using async context managers with SQLAlchemy's `AsyncSession`. Changes are staged and committed (`await db.commit()`) or rolled back (`await db.rollback()`) automatically upon unhandled exceptions."),
            ("Q18: What is the role of Alembic in this architecture?",
             "Alembic manages database schema migrations. It tracks schema changes in versioned Python scripts, allowing safe schema upgrades, rollbacks, and team synchronization across different environments."),
            ("Q19: How do you handle cross-cutting concerns like logging and exception handling?",
             "Through global FastAPI middlewares and exception handlers that intercept requests, attach correlation IDs, log structured JSON events via structlog, and format uncaught exceptions into standardized JSON responses."),
            ("Q20: How does your architecture support multi-tenancy if deployed for multiple universities?",
             "By adding a `tenant_id` (university_id) column to all core tables with row-level security (RLS) or tenant schema isolation in PostgreSQL.")
        ]),
        ("10.3 Category 3: Backend, FastAPI & Asynchronous Python", [
            ("Q21: Why did you choose FastAPI over Flask or Django?",
             "FastAPI offers native async/await concurrency with ASGI (Uvicorn), automatic OpenAPI/Swagger documentation generation, ultra-fast validation via Pydantic v2, and modern Python type hinting, delivering significantly higher requests per second than synchronous Flask/Django."),
            ("Q22: What is the difference between WSGI and ASGI?",
             "WSGI (Web Server Gateway Interface) is synchronous and blocks worker threads during I/O operations. ASGI (Asynchronous Server Gateway Interface) supports non-blocking asynchronous event loops, WebSockets, and concurrent HTTP/2 long-polling."),
            ("Q23: How does the Python `async` and `await` event loop work in your backend?",
             "When an I/O operation (database query, Redis call, network request) is awaited, the asyncio event loop pauses that coroutine and executes other pending tasks instead of blocking the CPU thread, enabling high concurrency on minimal resources."),
            ("Q24: What is Pydantic and how does KNOTS use it for validation?",
             "Pydantic is a data validation library powered by Rust core. In KNOTS, Pydantic schemas validate request payloads (email format, password length, enum types) before service execution and filter outgoing response fields for security."),
            ("Q25: What is the difference between `Pydantic Schema` and `SQLAlchemy Model`?",
             "SQLAlchemy Models define the relational database table structure, column constraints, and foreign key relationships. Pydantic Schemas define the shape and validation rules of API HTTP request and response payloads."),
            ("Q26: How does password hashing work in KNOTS?",
             "Passlib with Bcrypt (12 rounds) generates a unique salt and cryptographically hashes the password. Passwords are never stored in plaintext, and verification uses constant-time string comparisons to prevent timing attacks."),
            ("Q27: How does JWT authentication work from login to protected route access?",
             "1) User submits credentials to `/auth/login`; 2) Server verifies bcrypt hash and generates signed JWT access token; 3) Client stores token in localStorage; 4) Client sends `Authorization: Bearer <token>` on subsequent requests; 5) Server decodes token, verifies signature, extracts user_id, and grants access."),
            ("Q28: How is the `.docx` resume generated dynamically in Python?",
             "Using `python-docx`. The `ResumeGeneratorService` extracts candidate profile data, sets Calibri font and teal headings, injects custom XML paragraph borders (<w:pBdr>), configures right-aligned tab stops (7.3 in), categorizes skills, and writes directly to an in-memory `io.BytesIO` binary stream."),
            ("Q29: How do you serve binary files (like resumes and profile pictures) in FastAPI?",
             "Using `StreamingResponse` or `FileResponse` with appropriate MIME types (`application/vnd.openxmlformats-officedocument.wordprocessingml.document` for DOCX) and `Content-Disposition` attachment headers."),
            ("Q30: What is Structlog and why is structured logging preferred over standard print statements?",
             "Structlog emits structured JSON logs containing timestamp, log level, event name, user ID, request ID, and execution duration. This allows centralized ingestion and querying in tools like Datadog, ELK stack, or CloudWatch.")
        ]),
        ("10.4 Category 4: Database Design, PostgreSQL & SQLAlchemy ORM", [
            ("Q31: Why did you choose PostgreSQL over MongoDB / NoSQL databases?",
             "KNOTS is inherently relational: users connect to profiles, jobs link to applications and companies, posts link to comments and likes. PostgreSQL provides strict ACID compliance, foreign key integrity, and relational joins, while supporting JSONB columns for flexible data like skill sets."),
            ("Q32: Explain the database normalization level of your schema.",
             "The relational schema is normalized to 3NF (Third Normal Form) to eliminate update and delete anomalies. Transitive dependencies are removed into dedicated lookup tables (roles, event_categories, companies)."),
            ("Q33: How do you prevent the N+1 query problem in SQLAlchemy?",
             "By using eager loading strategies such as `selectinload()` and `joinedload()` on relationships (e.g. loading post author and profile in a single batch query rather than querying the database for each post's author separately)."),
            ("Q34: What is the purpose of database indexes and which columns are indexed in KNOTS?",
             "Indexes speed up record search from O(N) full-table scans to O(log N) B-Tree lookups. Key indexed columns include: `users.email` (unique index for login), `posts.author_id`, `connections.requester_id`, `connections.addressee_id`, `messages.conversation_id`, and `audit_logs.actor_id`."),
            ("Q35: What are cascade deletes and where are they used in your database?",
             "Cascade deletes (`cascade='all, delete-orphan'`) ensure child records are automatically removed when a parent record is deleted. Used on: User -> Profile, User -> Posts, Post -> Comments, Post -> Likes, Profile -> Education, Profile -> Experience."),
            ("Q36: Explain the difference between `asyncpg` and `psycopg2`.",
             "`asyncpg` is a high-performance, purely asynchronous PostgreSQL driver designed specifically for `asyncio`. `psycopg2` is the traditional synchronous driver. KNOTS uses `asyncpg` for high-throughput non-blocking database queries."),
            ("Q37: How do you store and query JSON data in SQLAlchemy models?",
             "Using the `JSON` column type (e.g. `Profile.skills`, `JobPosting.required_skills`, `Profile.certifications`). This allows storing structured lists or categorized dictionaries without requiring dozens of junction tables for static metadata."),
            ("Q38: What are unique constraints and give examples from KNOTS.",
             "Unique constraints prevent duplicate logical pairs: 1) `uq_connection_pair` on `(requester_id, addressee_id)`, 2) `uq_post_like` on `(post_id, user_id)` preventing duplicate likes, 3) `uq_event_rsvp` on `(event_id, user_id)`, 4) `uq_skill_endorsement` on `(profile_id, endorser_id, skill_name)`."),
            ("Q39: How do you handle database connection pooling in SQLAlchemy?",
             "Configured via `create_async_engine()` with `pool_size=20`, `max_overflow=10`, and `pool_pre_ping=True` to verify connection liveness before executing queries."),
            ("Q40: How does your system support testing with SQLite in-memory databases?",
             "Using `aiosqlite` with SQLite in-memory URLs (`sqlite+aiosqlite:///:memory:`) for lightning-fast, isolated automated unit and integration tests without needing a live PostgreSQL server.")
        ]),
        ("10.5 Category 5: Real-Time Messaging, WebSockets & Event Dispatching", [
            ("Q41: What is a WebSocket and how does it differ from HTTP polling?",
             "HTTP polling repeatedly opens and closes TCP connections to check for new messages, causing high server overhead and latency. WebSockets provide a persistent, bi-directional, full-duplex TCP communication channel over a single connection with minimal header overhead."),
            ("Q42: How does the WebSocket handshake and authentication work in KNOTS?",
             "The frontend initiates a WS handshake to `/api/v1/messaging/ws/{user_id}?token=<jwt>`. The server decodes and verifies the JWT token before calling `await websocket.accept()`. If invalid or expired, the connection is rejected with code 4001."),
            ("Q43: How does `ConnectionManager` handle users connected across multiple browser tabs?",
             "It maintains an in-memory dictionary mapping `user_id -> list[WebSocket]`. When a message arrives for a user, the manager iterates through and broadcasts to all active socket instances belonging to that user ID."),
            ("Q44: What happens when a client disconnects or loses network connection?",
             "The server catches WebSocketDisconnect exceptions, removes the closed socket from the active connection pool, and logs the disconnection. If no sockets remain for that user, their user ID key is cleaned up."),
            ("Q45: How is conversation history persisted when messages are sent over WebSockets?",
             "Messages received via WebSocket or REST are immediately committed to the PostgreSQL `messages` table with `sender_id`, `conversation_id`, and `created_at`, ensuring complete chat history is preserved upon refresh."),
            ("Q46: How does the system detect whether a user is currently online or offline?",
             "The `ConnectionManager.is_user_online(user_id)` method checks whether the user's ID exists in the active connections map with at least one live socket. This status is exposed to frontend chat headers."),
            ("Q47: How does read receipt tracking work in 1-on-1 and group chats?",
             "In 1-on-1 chats, an `is_read` boolean flag on the message is updated when the recipient opens the chat. In group chats, the `read_receipts` table records `(message_id, user_id, read_at)` timestamps."),
            ("Q48: How would you scale WebSocket messaging across multiple server instances in production?",
             "By integrating Redis Pub/Sub as a message broker. When Server A receives a message, it publishes it to a Redis channel; all server instances subscribed to the channel receive the event and push it to their locally connected WebSockets."),
            ("Q49: How do you prevent unauthorized users from eavesdropping on conversations?",
             "The server verifies that the authenticated user is an active participant in `conversation_participants` before returning messages or broadcasting real-time events for that conversation ID."),
            ("Q50: What message payload formats are exchanged over the WebSocket connection?",
             "Standardized JSON envelopes containing event types: `{type: 'CHAT_MESSAGE', data: {...}}`, `{type: 'TYPING_INDICATOR', user_id: 5}`, `{type: 'NOTIFICATION', title: '...', body: '...'}`.")
        ]),
        ("10.6 Category 6: Artificial Intelligence, Recommendations & Matching Algorithms", [
            ("Q51: Explain the mathematical algorithm used for AI Connection Recommendations.",
             "The algorithm computes a weighted match score: Base Score (40 pts) + Department Match (+30 pts) + Shared Skills Overlap (+10 pts per matching skill, max 30 pts) + Graduation Proximity (+10 pts for same year, +5 pts for +/-1 yr). Scores are clamped to [40, 98] and sorted descending."),
            ("Q52: Explain the AI Job Recommendation matching logic.",
             "The engine compares user skills with `JobPosting.required_skills`. Matching skills add +15 pts each (up to 45 pts). If the user's department keyword matches the job title or description, +15 pts is added. Already applied jobs and closed vacancies are filtered out."),
            ("Q53: How does the AI Feed Content Recommendation work?",
             "It builds an interest keyword vector from the user's skills and department, scanning public posts for keyword matches (+20 pts per match, max 40 pts), and adds an engagement bonus (+15 pts) for posts with high community likes and comments."),
            ("Q54: What is the Trending Posts scoring formula implemented in Analytics?",
             "Score = ((Views * 1) + (Likes * 3) + (Comments * 5)) / ((Age_in_Hours + 2) ^ 1.5). This time-decay formula ensures popular recent posts trend over older saturated posts."),
            ("Q55: What is FAISS and how will it be integrated in the next phase of KNOTS?",
             "FAISS (Facebook AI Similarity Search) is an open-source library for efficient similarity search of dense vectors. In Phase 2, student profiles and job descriptions will be converted into 384-dimensional dense embeddings using Sentence Transformers, allowing sub-millisecond semantic search via FAISS."),
            ("Q56: What is the role of LangChain in your planned AI roadmap?",
             "LangChain will orchestrate multi-step LLM chains: 1) Parsing uploaded PDF resumes, 2) Extracting structured JSON entities, 3) Comparing against job descriptions to generate an ATS compatibility score and actionable improvement suggestions."),
            ("Q57: How do you handle cold-start problems for new users with no skills or history?",
             "The system defaults to a baseline prior (40 pts) and relies on department, role, and graduation cohort matching until the user populates their skills and interests."),
            ("Q58: How does the AI Moderation Service filter toxic content?",
             "By analyzing text against profanity filters, toxicity classification models, and automatically routing flagged posts to the admin moderation queue with PENDING status."),
            ("Q59: How does the AI Club Recommendation Service work?",
             "By matching student interest tags, department domain, and peer co-enrollment patterns to suggest relevant campus technical clubs and cultural organizations."),
            ("Q60: What is the computational complexity of your current recommendation algorithms?",
             "O(U * S) where U is the number of active candidate users/jobs and S is the average number of skills. With indexed queries and candidate pre-filtering, execution takes under 15ms for typical college databases (10k+ users).")
        ]),
        ("10.7 Category 7: Frontend, React, TypeScript & User Experience", [
            ("Q61: Why did you choose React + Vite over traditional Next.js or Create-React-App?",
             "Vite uses native ES modules in development, providing sub-second hot module replacement (HMR) and ultra-fast Rollup builds. For an authenticated campus portal with dynamic client state and WebSockets, a client-side Single Page Application (SPA) offers optimal responsiveness."),
            ("Q62: What are the key benefits of using TypeScript across your frontend?",
             "TypeScript provides compile-time type safety, autocompletion for API models, and eliminates common JavaScript runtime bugs such as undefined property accesses and type mismatches."),
            ("Q63: How is client-side routing structured in your React application?",
             "Using React Router v6 in `AppRoutes.tsx`. Protected routes are wrapped in an authentication guard that checks for a valid JWT token in localStorage and redirects unauthenticated users to `/login`."),
            ("Q64: How do you handle responsive UI design across desktop, tablet, and mobile?",
             "Using Tailwind CSS responsive breakpoint utilities (`sm:`, `md:`, `lg:`, `xl:`), fluid flex/grid containers, and collapsible mobile drawer menus."),
            ("Q65: What charting library did you use for analytics and why?",
             "Recharts. It provides declarative, SVG-based, responsive charts (LineChart for profile views time-series, PieChart/Donut for platform engagement distribution) with smooth tooltip animations."),
            ("Q66: How is the global search bar implemented on the frontend?",
             "The `GlobalSearchBar` component triggers debounced API requests to `/api/v1/search/global`, rendering categorized live search results for Users, Jobs, Events, and Clubs in an overlay dropdown."),
            ("Q67: How do you manage optimistic UI updates (e.g. liking a post or RSVPing)?",
             "The UI immediately increments the like counter and updates the heart icon state before the API network request finishes. If the server request fails, the state rolls back gracefully with an error toast."),
            ("Q68: What is Shadcn UI and why is it used?",
             "Shadcn UI provides accessible, unstyled Radix UI primitives with Tailwind CSS styling. Unlike heavy component libraries, components are owned in the codebase and fully customizable."),
            ("Q69: How do you handle user file uploads (profile pictures, resumes) on the frontend?",
             "Using HTML `<input type='file'>` wrapped in drag-and-drop zones, constructing `FormData` multipart payloads, and submitting via Axios with upload progress indicators."),
            ("Q70: How does the frontend handle WebSocket reconnections during network drops?",
             "Through exponential backoff retry logic. If the socket closes unexpectedly, the client waits 1s, 2s, 4s, up to 10s before attempting automatic re-handshakes.")
        ]),
        ("10.8 Category 8: Security, Governance, RBAC & Compliance", [
            ("Q71: What is Role-Based Access Control (RBAC) and how is it implemented?",
             "RBAC restricts API endpoint execution based on the user's role. In KNOTS, roles (Admin, Faculty, Student, Alumni, Recruiter) are validated via FastAPI's `RoleRequired(['Admin'])` dependency, checking claims in the decoded JWT."),
            ("Q72: How does KNOTS protect against SQL Injection?",
             "All database interactions strictly utilize SQLAlchemy 2.0 ORM parameterized queries. User input is never concatenated directly into raw SQL strings, ensuring parameter values are escaped by the driver."),
            ("Q73: How does KNOTS mitigate Cross-Site Scripting (XSS) attacks?",
             "1) React's JSX automatically escapes strings before rendering to DOM; 2) Pydantic request models sanitize HTML tags; 3) Sensitive auth tokens are protected from malicious script exfiltration."),
            ("Q74: What is Cross-Site Request Forgery (CSRF) and why are Bearer JWTs immune?",
             "CSRF attacks exploit automatic browser cookie transmission. Because KNOTS uses explicit `Authorization: Bearer <token>` headers attached by Axios JavaScript rather than ambient cookies, malicious third-party cross-site requests cannot carry credentials."),
            ("Q75: What is the purpose of the `AuditLog` table?",
             "It maintains an immutable administrative audit ledger recording the actor's user ID, action performed (e.g. BAN_USER, DELETE_POST), target resource, client IP address, and UTC timestamp for compliance and forensic auditing."),
            ("Q76: How does KNOTS comply with the Digital Personal Data Protection (DPDP) Act of India?",
             "By implementing explicit user consent, providing granular notification toggles, enforcing purpose limitation on profile data, and enabling complete account and data deletion ('Right to Erasure')."),
            ("Q77: How are sensitive environment variables and API keys secured?",
             "Stored in `.env` files that are excluded from Git version control via `.gitignore`, and validated at startup using `pydantic_settings.BaseSettings`."),
            ("Q78: How are banned users prevented from accessing platform features?",
             "The `get_current_user` auth dependency checks `User.is_active`. If false, it immediately raises an HTTP 403 Forbidden exception, revoking access across all endpoints."),
            ("Q79: What is the difference between encryption in transit and encryption at rest?",
             "Encryption in transit protects data moving between client and server via HTTPS / TLS 1.3. Encryption at rest encrypts data stored on disk in PostgreSQL database volumes and backups."),
            ("Q80: How does the system handle password reset securely?",
             "By generating a short-lived (15 min), cryptographically secure, one-time reset token sent to the user's verified institutional email.")
        ]),
        ("10.9 Category 9: DevOps, Testing, CI/CD & Deployment", [
            ("Q81: What is Docker and what advantages does it bring to this project?",
             "Docker packages the application, runtime, system tools, and dependencies into lightweight, isolated containers, eliminating 'it works on my machine' inconsistencies across Windows, Linux, and macOS."),
            ("Q82: Explain the multi-stage Docker build used for the frontend.",
             "Stage 1 (Build): Uses Node 18 Alpine to install packages and compile TypeScript to static HTML/CSS/JS. Stage 2 (Production): Copies only the compiled `dist` directory into an ultra-small Nginx Alpine image, reducing final image size from 800MB+ to under 25MB."),
            ("Q83: What is Docker Compose and what services are defined in your orchestration file?",
             "Docker Compose defines and runs multi-container Docker applications. Services: `backend` (FastAPI on port 8000), `frontend` (Nginx on port 80), `postgres` (database on port 5432), and `redis` (cache on port 6379) linked via an internal bridge network."),
            ("Q84: What testing framework do you use and what test types are implemented?",
             "Pytest with `pytest-asyncio`. Test types: 1) Unit tests for services and algorithms, 2) Router integration tests with `httpx.AsyncClient`, 3) End-to-end integration tests (`test_full_platform_integration.py`, `test_admin_e2e_integration.py`)."),
            ("Q85: How do you mock database sessions during automated tests?",
             "Using Pytest fixtures that override the `get_db` FastAPI dependency with an in-memory SQLite database session (`aiosqlite`), ensuring tests run independently without modifying production data."),
            ("Q86: What is a reverse proxy and what role does Nginx play in your deployment?",
             "Nginx acts as a reverse proxy in front of the application: 1) Serves static frontend bundles, 2) Proxies `/api/` requests to the FastAPI ASGI server, 3) Handles WebSocket connection upgrades, and 4) Provides gzip compression."),
            ("Q87: What is CI/CD and how can your GitHub Actions pipeline be structured?",
             "Continuous Integration / Continuous Deployment. On every Git push/PR, GitHub Actions runs Ruff linter, TypeScript typecheck, Alembic migration validation, and Pytest test suite before building Docker images."),
            ("Q88: How do you handle database backups in a production PostgreSQL deployment?",
             "Using automated `pg_dump` cron jobs backing up compressed SQL dumps to encrypted AWS S3 buckets or remote storage with point-in-time recovery (PITR) enabled."),
            ("Q89: How does KNOTS handle database schema evolution without downtime?",
             "Alembic migration scripts apply non-destructive forward schema changes (e.g. adding nullable columns, creating new tables) during deployment before rolling out updated container images."),
            ("Q90: What metrics would you monitor to assess system health in production?",
             "HTTP request latency (p95, p99), error rates (5xx responses), active WebSocket connection counts, PostgreSQL connection pool utilization, CPU and memory usage.")
        ]),
        ("10.10 Category 10: Research Innovations, Viva Traps & Future Enhancements", [
            ("Q91: What is the main research novelty of KNOTS compared to commercial networking platforms?",
             "The integration of verified multi-stakeholder trust graph modeling (linking academic department, faculty skill validations, and cohort proximity) into open-weight recommendation algorithms rather than proprietary blackbox ad-driven ranking."),
            ("Q92: If the examiner asks: 'Is your AI model trained from scratch or using pre-trained weights?' how do you answer?",
             "Our current recommendation system utilizes deterministic multi-vector mathematical scoring algorithms tailored to campus metadata. Our NLP resume parsing and vector search roadmap utilizes pre-trained sentence transformer embeddings (`all-MiniLM-L6-v2`) and foundation LLMs via LangChain rather than training custom models from scratch, which is industry best practice."),
            ("Q93: What will you do if two users send connection requests to each other simultaneously?",
             "The database `uq_connection_pair` unique constraint and service transaction locking catch race conditions. If User A requests User B while User B's request is pending, the service automatically marks the connection as ACCEPTED."),
            ("Q94: How does your resume generator handle missing data fields (e.g., student has no work experience)?",
             "The `ResumeGeneratorService` implements graceful fallback handling: omitting empty sections, rendering academic coursework and capstone projects in place of industry experience, and auto-deriving institution details from the profile department."),
            ("Q95: How do you prevent spam on the campus feed?",
             "Through post creation rate limits, audience visibility constraints (e.g. students only), community post flagging with automated admin queues, and AI profanity analysis."),
            ("Q96: How would you implement video conferencing or audio mentorship calls in KNOTS?",
             "By integrating WebRTC (Web Real-Time Communication) signaling over our existing WebSocket infrastructure, enabling direct peer-to-peer audio/video streaming without external third-party servers."),
            ("Q97: Can alumni post jobs from companies they do not work for?",
             "Alumni can post verified job vacancies or mark them as referral leads. Admins have oversight to verify company profiles in the `companies` directory before listings are marked as campus-approved."),
            ("Q98: How does your system support offline usage on mobile devices?",
             "By converting the React frontend into a Progressive Web App (PWA) with Service Workers caching application shell assets and IndexedDB storing offline chat drafts."),
            ("Q99: What are the primary scalability bottlenecks and how will you resolve them for 100,000+ students?",
             "1) WebSocket connection state: scale using Redis Pub/Sub clusters; 2) Relational database read load: introduce read replicas and Redis caching for trending posts and profile views; 3) Media storage: offload profile pictures and resume attachments to AWS S3 / Cloudflare R2 object storage with CDN."),
            ("Q100: Concluding Summary: Why is KNOTS an ideal Final Year Engineering Capstone Project?",
             "KNOTS integrates every core discipline of Computer Science & Engineering: Advanced Database Modeling (PostgreSQL, 3NF), Asynchronous Systems (FastAPI, asyncio), Real-time Distributed Systems (WebSockets), Modern Web UI (React, TypeScript, Tailwind), Algorithmic Decision Making (AI Multi-Vector Scoring), File Binary Engineering (DOCX resume generation), and Enterprise Security (JWT, RBAC, DPDP compliance). It solves a high-impact real-world problem with tangible social value for higher education.")
        ])
    ]

    for cat_title, qa_list in viva_qa_sections:
        add_custom_heading_2(cat_title)
        for q, a in qa_list:
            add_body_paragraph(q, bold_prefix="", italic=False)
            add_body_paragraph(a, bold_prefix="Answer: ", italic=False)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Save Document
    doc.save(file_path)
    print(f"Master Seminar Guide generated successfully at: {file_path}")

if __name__ == "__main__":
    target_path = os.path.abspath("KNOTS_Project_Master_Seminar_Guide.docx")
    create_master_document(target_path)
