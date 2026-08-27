import io
from datetime import date
from typing import Any
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

PRIMARY_COLOR = RGBColor(0, 136, 168)  # Cyan/Teal heading color from template (#0088A8)
TEXT_DARK = RGBColor(20, 24, 33)  # Deep black/navy text color (#141821)
TEXT_MUTED = RGBColor(70, 80, 95)  # Muted subtext color


def format_date_range(start_date: Any, end_date: Any) -> str:
    """Format start and end dates into 'Mon YYYY - Mon YYYY' or 'Present'."""

    def _format_single(d: Any) -> str:
        if not d:
            return ""
        if isinstance(d, date):
            return d.strftime("%b %Y")
        if isinstance(d, str):
            try:
                parts = d.split("-")
                if len(parts) >= 2:
                    year, month = int(parts[0]), int(parts[1])
                    temp = date(year, month, 1)
                    return temp.strftime("%b %Y")
            except Exception:
                return str(d)
        return str(d)

    start_str = _format_single(start_date)
    end_str = _format_single(end_date) if end_date else "Present"

    if start_str and end_str:
        return f"{start_str} - {end_str}"
    if start_str:
        return start_str
    if end_str and end_str != "Present":
        return end_str
    return ""


def add_section_header(doc: docx.Document, title: str):
    """Add an uppercase, bold, teal section header with a solid black bottom rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True

    run = p.add_run(title.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    # Bottom border rule in XML
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="2" w:color="000000"/>'
        f"</w:pBdr>"
    )
    pPr.append(pBdr)


def add_two_column_line(
    doc: docx.Document,
    left_text: str,
    right_text: str,
    is_bold_left: bool = False,
    is_bold_right: bool = False,
    space_before: int = 4,
    space_after: int = 1,
):
    """Add a line with left-aligned text and right-aligned text across page width."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True

    # 7.3 inches text width (8.5 - 2*0.6)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.3), WD_TAB_ALIGNMENT.RIGHT)

    left_run = p.add_run(left_text)
    left_run.font.name = "Calibri"
    left_run.font.size = Pt(10)
    left_run.font.bold = is_bold_left
    left_run.font.color.rgb = TEXT_DARK

    if right_text:
        p.add_run("\t")
        right_run = p.add_run(right_text)
        right_run.font.name = "Calibri"
        right_run.font.size = Pt(10)
        right_run.font.bold = is_bold_right
        right_run.font.color.rgb = TEXT_DARK


def add_bullet_point(doc: docx.Document, text: str, prefix_bold: str = ""):
    """Add a clean, tight bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Inches(0.25)

    if prefix_bold:
        b_run = p.add_run(prefix_bold)
        b_run.font.name = "Calibri"
        b_run.font.size = Pt(9.5)
        b_run.font.bold = True
        b_run.font.color.rgb = TEXT_DARK

    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.font.color.rgb = TEXT_DARK


class ResumeGeneratorService:
    """Generates a professional DOCX resume matching the specified clean tech template."""

    @staticmethod
    def generate_docx(profile_data: dict, user_email: str = "") -> io.BytesIO:
        doc = docx.Document()

        # Set 0.6 inch margins for modern single/dual page layout
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)

        # Set default font
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(10)

        # Extract profile details
        first_name = (profile_data.get("first_name") or "").strip()
        last_name = (profile_data.get("last_name") or "").strip()
        full_name = f"{first_name} {last_name}".strip()
        if not full_name:
            email_handle = (
                user_email.split("@")[0] if user_email else "Student Candidate"
            )
            parts = [
                p.capitalize() for p in email_handle.replace("_", ".").split(".") if p
            ]
            full_name = " ".join(parts) if parts else "Student Candidate"

        department = profile_data.get("department") or "Computer Science & Engineering"
        grad_year = profile_data.get("graduation_year")

        # -------------------------------------------------------------
        # 1. HEADER: Candidate Name & Contact Info
        # -------------------------------------------------------------
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_p.paragraph_format.space_before = Pt(0)
        name_p.paragraph_format.space_after = Pt(2)
        name_run = name_p.add_run(full_name)
        name_run.font.name = "Calibri"
        name_run.font.size = Pt(20)
        name_run.font.bold = True
        name_run.font.color.rgb = TEXT_DARK

        # Contact line
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_before = Pt(0)
        contact_p.paragraph_format.space_after = Pt(6)

        clean_handle = "".join(e for e in full_name.lower() if e.isalnum())
        if not clean_handle:
            clean_handle = (
                user_email.split("@")[0].lower() if user_email else "candidate"
            )

        contact_items = []
        if user_email:
            contact_items.append(f"✉  {user_email}")
        contact_items.append("📞  +91 9876543210")
        contact_items.append(f"in  {clean_handle}")
        contact_items.append(f"🐙  {clean_handle}")

        contact_run = contact_p.add_run("    |    ".join(contact_items))
        contact_run.font.name = "Calibri"
        contact_run.font.size = Pt(9.5)
        contact_run.font.color.rgb = TEXT_MUTED

        # -------------------------------------------------------------
        # 2. WORK EXPERIENCE
        # -------------------------------------------------------------
        employment = profile_data.get("employment_history") or []
        if employment:
            add_section_header(doc, "Work Experience")
            for emp in employment:
                company = emp.get("company_name", "").strip() or "Tech Company"
                title = emp.get("title", "").strip() or "Software Engineer"
                date_str = format_date_range(emp.get("start_date"), emp.get("end_date"))
                location = emp.get("location", "").strip()

                # Line 1: Company Name (Bold)
                add_two_column_line(
                    doc,
                    company,
                    location,
                    is_bold_left=True,
                    space_before=4,
                    space_after=1,
                )
                # Line 2: Title & Dates
                add_two_column_line(
                    doc,
                    title,
                    date_str,
                    is_bold_left=False,
                    space_before=0,
                    space_after=2,
                )

                desc = emp.get("description") or ""
                if desc:
                    lines = [
                        line.strip().lstrip("•-*").strip()
                        for line in desc.split("\n")
                        if line.strip()
                    ]
                    for line_item in lines:
                        add_bullet_point(doc, line_item)
                else:
                    add_bullet_point(
                        doc,
                        f"Contributed to core development, system architecture, and feature delivery for {company}.",
                    )
                    add_bullet_point(
                        doc,
                        "Collaborated with cross-functional teams to design scalable solutions and improve user experience.",
                    )

        # -------------------------------------------------------------
        # 3. EDUCATION
        # -------------------------------------------------------------
        education_list = profile_data.get("education") or []
        add_section_header(doc, "Education")
        if education_list:
            for edu in education_list:
                inst = (
                    edu.get("institution_name", "").strip()
                    or "S.B. Jain Institute of Technology, Management & Research"
                )
                degree = edu.get("degree", "").strip() or "Bachelor of Technology"
                field = edu.get("field_of_study", "").strip()
                gpa = edu.get("gpa")
                date_str = format_date_range(edu.get("start_date"), edu.get("end_date"))

                degree_full = f"{degree} in {field}" if field else degree
                if gpa:
                    degree_full += f" - {gpa} GPA"

                # Line 1: Institution (Bold)
                add_two_column_line(
                    doc, inst, "", is_bold_left=True, space_before=4, space_after=1
                )
                # Line 2: Degree, GPA, Date Range
                add_two_column_line(
                    doc,
                    degree_full,
                    date_str,
                    is_bold_left=False,
                    space_before=0,
                    space_after=2,
                )

                desc = edu.get("description") or ""
                if desc:
                    lines = [
                        line.strip().lstrip("•-*").strip()
                        for line in desc.split("\n")
                        if line.strip()
                    ]
                    for line_item in lines:
                        add_bullet_point(doc, line_item)
        else:
            # Fallback based on profile department and grad year
            inst = "S.B. Jain Institute of Technology, Management & Research, Nagpur"
            degree_full = f"Bachelor of Technology in {department}"
            grad_str = f"Class of {grad_year}" if grad_year else "2023 - 2027"
            add_two_column_line(
                doc, inst, "", is_bold_left=True, space_before=4, space_after=1
            )
            add_two_column_line(
                doc,
                degree_full,
                grad_str,
                is_bold_left=False,
                space_before=0,
                space_after=2,
            )
            add_bullet_point(
                doc,
                "Relevant Coursework: Data Structures, Object-Oriented Programming, Database Systems, Web Technologies.",
            )

        # -------------------------------------------------------------
        # 4. PROJECTS
        # -------------------------------------------------------------
        projects = profile_data.get("projects") or []
        if projects:
            add_section_header(doc, "Project")
            for proj in projects:
                p_title = (
                    proj.get("title") or proj.get("name") or "Key Software Project"
                )
                tech_stack = proj.get("tech_stack") or []
                highlights = proj.get("highlights") or []
                p_desc = proj.get("description") or ""

                # Project Title (Bold)
                p_para = doc.add_paragraph()
                p_para.paragraph_format.space_before = Pt(4)
                p_para.paragraph_format.space_after = Pt(1)
                p_para.paragraph_format.keep_with_next = True
                p_run = p_para.add_run(p_title)
                p_run.font.name = "Calibri"
                p_run.font.size = Pt(10)
                p_run.font.bold = True
                p_run.font.color.rgb = TEXT_DARK

                if highlights:
                    for h in highlights:
                        add_bullet_point(doc, h)
                elif p_desc:
                    lines = [
                        line.strip().lstrip("•-*").strip()
                        for line in p_desc.split("\n")
                        if line.strip()
                    ]
                    for line_item in lines:
                        add_bullet_point(doc, line_item)
                else:
                    tech_str = (
                        ", ".join(tech_stack) if tech_stack else "modern technologies"
                    )
                    add_bullet_point(
                        doc,
                        f"Designed and developed full-stack application using {tech_str}.",
                    )
                    add_bullet_point(
                        doc,
                        "Implemented secure user authentication, responsive UI, and optimized database queries.",
                    )

                if tech_stack and not highlights and not p_desc:
                    add_bullet_point(
                        doc, f"Technologies used: {', '.join(tech_stack)}."
                    )

        # -------------------------------------------------------------
        # 5. SKILLS
        # -------------------------------------------------------------
        raw_skills = profile_data.get("skills")
        add_section_header(doc, "Skills")

        if isinstance(raw_skills, dict) and raw_skills:
            for category, skill_list in raw_skills.items():
                if isinstance(skill_list, list) and skill_list:
                    cat_name = category.strip()
                    skills_str = ", ".join(
                        str(s).strip() for s in skill_list if str(s).strip()
                    )
                    add_bullet_point(doc, skills_str, prefix_bold=f"{cat_name}: ")
        elif isinstance(raw_skills, list) and raw_skills:
            # Categorize flat list if possible or output categorized bullets
            all_skills = [str(s).strip() for s in raw_skills if str(s).strip()]
            frontend_kw = {
                "html",
                "css",
                "javascript",
                "typescript",
                "react",
                "react.js",
                "angular",
                "vue",
                "tailwind",
                "next.js",
                "bootstrap",
            }
            backend_kw = {
                "node",
                "node.js",
                "express",
                "express.js",
                "python",
                "fastapi",
                "django",
                "java",
                "spring",
                "spring boot",
                "c++",
                "c#",
                "rest",
                "graphql",
            }
            db_tools_kw = {
                "sql",
                "postgresql",
                "mysql",
                "mongodb",
                "redis",
                "git",
                "github",
                "docker",
                "aws",
                "kubernetes",
                "linux",
                "ci/cd",
            }

            fe, be, tools, other = [], [], [], []
            for s in all_skills:
                sl = s.lower()
                if any(k in sl for k in frontend_kw):
                    fe.append(s)
                elif any(k in sl for k in backend_kw):
                    be.append(s)
                elif any(k in sl for k in db_tools_kw):
                    tools.append(s)
                else:
                    other.append(s)

            if fe:
                add_bullet_point(doc, ", ".join(fe) + ".", prefix_bold="Front-end: ")
            if be:
                add_bullet_point(doc, ", ".join(be) + ".", prefix_bold="Back-end: ")
            if tools:
                add_bullet_point(
                    doc, ", ".join(tools) + ".", prefix_bold="Tools & Databases: "
                )
            if other or (not fe and not be and not tools):
                add_bullet_point(
                    doc,
                    ", ".join(other or all_skills) + ".",
                    prefix_bold="Core Skills: ",
                )
        else:
            add_bullet_point(
                doc,
                "HTML, CSS, JavaScript, TypeScript, React.js, and Responsive Design.",
                prefix_bold="Front-end: ",
            )
            add_bullet_point(
                doc,
                "Python, FastAPI, Node.js, Express.js, and RESTful API design.",
                prefix_bold="Back-end: ",
            )
            add_bullet_point(
                doc,
                "Git, GitHub, PostgreSQL, Docker, and CI/CD pipelines.",
                prefix_bold="Tools: ",
            )

        # -------------------------------------------------------------
        # 6. CERTIFICATIONS
        # -------------------------------------------------------------
        certs = profile_data.get("certifications") or []
        if certs:
            add_section_header(doc, "Certifications")
            for cert in certs:
                if isinstance(cert, dict):
                    c_name = (
                        cert.get("name")
                        or cert.get("title")
                        or "Professional Certification"
                    )
                    c_issuer = cert.get("issuer")
                    c_text = f"{c_name} - {c_issuer}" if c_issuer else c_name
                    add_bullet_point(doc, c_text)
                elif isinstance(cert, str):
                    add_bullet_point(doc, cert)

        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        return output_stream
